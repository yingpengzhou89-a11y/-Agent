import io
import re
import time
from collections import defaultdict
from pathlib import Path
from uuid import UUID, uuid4

import fitz
import httpx
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.errors import AppError
from app.models.knowledge import (
    DocumentChunk,
    KnowledgeDocument,
    KnowledgeSearchFeedback,
)
from app.repositories.knowledge import KnowledgeRepository
from app.schemas.knowledge import (
    BulkReindexRead,
    EmbeddingStatusRead,
    KnowledgeRetrievalQualityRead,
    KnowledgeSearchResult,
)


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}


def extract_text(filename: str, raw: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise AppError("FILE_PARSE_ERROR", "仅支持 PDF、DOCX、Markdown 和 TXT 文件", status_code=422)
    if suffix in {".md", ".txt"}:
        return raw.decode("utf-8", errors="replace").strip()
    try:
        if suffix == ".docx":
            document = Document(io.BytesIO(raw))
            return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        with fitz.open(stream=raw, filetype="pdf") as pdf:
            return "\n".join(page.get_text() for page in pdf).strip()
    except Exception as exc:
        raise AppError("FILE_PARSE_ERROR", "文件无法解析，请确认它没有损坏或加密", status_code=422) from exc


def _split_long_unit(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    sentences = [item.strip() for item in re.split(r"(?<=[。！？；.!?;])", text) if item.strip()]
    units: list[str] = []
    buffer = ""
    for sentence in sentences:
        if len(sentence) > max_chars:
            if buffer:
                units.append(buffer)
                buffer = ""
            units.extend(sentence[index : index + max_chars] for index in range(0, len(sentence), max_chars))
        elif buffer and len(buffer) + len(sentence) > max_chars:
            units.append(buffer)
            buffer = sentence
        else:
            buffer += sentence
    if buffer:
        units.append(buffer)
    return units


def chunk_text(text: str, max_chars: int = 900, overlap: int = 120) -> list[tuple[str, dict]]:
    """Create focused RAG chunks while retaining section heading and local context."""
    lines = [line.rstrip() for line in text.splitlines() if line.strip()]
    chunks: list[tuple[str, dict]] = []
    heading = ""
    buffer = ""

    def emit() -> None:
        nonlocal buffer
        if buffer.strip():
            chunks.append((buffer.strip(), {"heading": heading}))
        buffer = ""

    for line in lines:
        if line.lstrip().startswith("#"):
            emit()
            heading = line.lstrip("# ").strip()
            buffer = f"# {heading}"
            continue
        for unit in _split_long_unit(line, max_chars):
            candidate = f"{buffer}\n{unit}".strip()
            if len(candidate) > max_chars and buffer:
                previous = buffer
                emit()
                buffer = f"{previous[-overlap:]}\n{unit}".strip()
            else:
                buffer = candidate
    emit()
    return chunks


def _tokens(text: str) -> list[str]:
    return re.findall(r"[a-z0-9+#.]+|[\u4e00-\u9fff]{2,}", text.casefold())


def _character_ngrams(text: str, width: int = 2) -> set[str]:
    """Build language-neutral character n-grams for the no-embedding fallback.

    PostgreSQL's built-in ``simple`` text-search configuration has no Chinese
    tokenizer. Character n-grams preserve partial overlap when a user swaps
    phrase order, while ASCII terms are still scored as whole tokens.
    """
    normalized = "".join(re.findall(r"[a-z0-9+#.]|[\u4e00-\u9fff]", text.casefold()))
    if len(normalized) < width:
        return {normalized} if normalized else set()
    return {normalized[index : index + width] for index in range(len(normalized) - width + 1)}


def rrf_fuse(rankings: list[list[UUID]], k: int = 60) -> dict[UUID, float]:
    scores: dict[UUID, float] = defaultdict(float)
    for ranking in rankings:
        for position, identifier in enumerate(ranking, start=1):
            scores[identifier] += 1 / (k + position)
    return dict(scores)


class EmbeddingGateway:
    async def embed(self, texts: list[str]) -> list[list[float]]:
        if not settings.embedding_base_url or not settings.embedding_api_key:
            raise AppError("EMBEDDING_NOT_CONFIGURED", "未配置 Embedding 服务", status_code=503)
        endpoint = f"{settings.embedding_base_url.rstrip('/')}/embeddings"
        async with httpx.AsyncClient(timeout=settings.llm_timeout_seconds, trust_env=False) as client:
            response = await client.post(
                endpoint,
                headers={"Authorization": f"Bearer {settings.embedding_api_key}"},
                json={"model": settings.embedding_model, "input": texts},
            )
            try:
                response.raise_for_status()
                vectors = [item["embedding"] for item in response.json()["data"]]
            except (httpx.HTTPError, KeyError, TypeError) as exc:
                raise AppError("RETRIEVAL_ERROR", "Embedding 服务调用失败", retryable=True, status_code=502) from exc
        if any(len(vector) != settings.embedding_dimensions for vector in vectors):
            raise AppError("RETRIEVAL_ERROR", "Embedding 维度与配置不一致", status_code=502)
        return vectors


class KnowledgeService:
    def __init__(self, embeddings: EmbeddingGateway | None = None) -> None:
        self.repo = KnowledgeRepository()
        self.embeddings = embeddings or EmbeddingGateway()

    async def ingest_text(
        self, session: AsyncSession, user_id: UUID, name: str, source_type: str, text: str, file_path: str | None = None
    ) -> KnowledgeDocument:
        if not text.strip():
            raise AppError("FILE_PARSE_ERROR", "文档没有可索引的文本内容", status_code=422)
        document = await self.repo.create_document(
            session,
            KnowledgeDocument(
                user_id=user_id,
                name=name,
                source_type=source_type,
                file_path=file_path,
                parse_status="COMPLETED",
                index_status="PENDING",
            ),
        )
        chunks = [
            DocumentChunk(document_id=document.id, user_id=user_id, content=content, metadata_json=metadata)
            for content, metadata in chunk_text(text)
        ]
        await self.repo.replace_chunks(session, document, chunks)
        document.index_status = "KEYWORD_READY"
        await session.flush()
        return document

    async def ingest_file(
        self, session: AsyncSession, user_id: UUID, source_type: str, filename: str, raw: bytes
    ) -> KnowledgeDocument:
        if len(raw) > settings.max_upload_mb * 1024 * 1024:
            raise AppError("FILE_PARSE_ERROR", "上传文件超过大小限制", status_code=413)
        text = extract_text(filename, raw)
        storage_dir = Path(settings.storage_dir) / str(user_id)
        storage_dir.mkdir(parents=True, exist_ok=True)
        file_path = storage_dir / f"{uuid4().hex}{Path(filename).suffix.lower()}"
        file_path.write_bytes(raw)
        return await self.ingest_text(
            session, user_id, filename, source_type, text, file_path=str(file_path)
        )

    async def delete(self, session: AsyncSession, user_id: UUID, document_id: UUID) -> None:
        document = await self.repo.get_document(session, user_id, document_id)
        if document.file_path:
            path = Path(document.file_path)
            if path.is_file():
                path.unlink()
        await session.delete(document)

    async def reindex(self, session: AsyncSession, user_id: UUID, document_id: UUID) -> KnowledgeDocument:
        document = await self.repo.get_document(session, user_id, document_id)
        chunks = await self.repo.chunks_for_document(session, document.id)
        if not chunks:
            raise AppError("RETRIEVAL_ERROR", "文档没有可索引的片段", status_code=409)
        document.index_status = "INDEXING"
        vectors = await self.embeddings.embed([chunk.content for chunk in chunks])
        for chunk, vector in zip(chunks, vectors, strict=True):
            chunk.embedding = vector
        document.index_status = "INDEXED"
        await session.flush()
        return document

    async def rechunk_and_reindex(self, session: AsyncSession, user_id: UUID, document_id: UUID) -> KnowledgeDocument:
        """Reparse the original file, then atomically replace old chunks after embeddings succeed."""
        document = await self.repo.get_document(session, user_id, document_id)
        if not document.file_path or not Path(document.file_path).is_file():
            raise AppError("RETRIEVAL_ERROR", "原始文件不存在，无法重新切分", status_code=409)
        raw = Path(document.file_path).read_bytes()
        text = extract_text(document.name, raw)
        chunks = [
            DocumentChunk(document_id=document.id, user_id=user_id, content=content, metadata_json=metadata)
            for content, metadata in chunk_text(text)
        ]
        if not chunks:
            raise AppError("RETRIEVAL_ERROR", "文档没有可索引的片段", status_code=409)
        document.index_status = "INDEXING"
        vectors = await self.embeddings.embed([chunk.content for chunk in chunks])
        for chunk, vector in zip(chunks, vectors, strict=True):
            chunk.embedding = vector
        await self.repo.replace_chunks(session, document, chunks)
        document.index_status = "INDEXED"
        await session.flush()
        return document

    async def embedding_status(self, session: AsyncSession, user_id: UUID) -> EmbeddingStatusRead:
        documents = await self.repo.list_documents(session, user_id)
        indexed = sum(document.index_status == "INDEXED" for document in documents)
        return EmbeddingStatusRead(
            configured=bool(settings.embedding_base_url and settings.embedding_api_key),
            model=settings.embedding_model,
            dimensions=settings.embedding_dimensions,
            indexed_document_count=indexed,
            pending_document_count=len(documents) - indexed,
        )

    async def reindex_all(self, session: AsyncSession, user_id: UUID) -> BulkReindexRead:
        if not settings.embedding_base_url or not settings.embedding_api_key:
            raise AppError("EMBEDDING_NOT_CONFIGURED", "请先配置 Embedding 服务", status_code=409)
        documents = await self.repo.list_documents(session, user_id)
        failures: list[dict[str, str]] = []
        indexed_count = 0
        for document in documents:
            try:
                await self.reindex(session, user_id, document.id)
                indexed_count += 1
            except AppError as exc:
                document.index_status = "KEYWORD_READY"
                failures.append({"document_id": str(document.id), "name": document.name, "reason": exc.message})
        await session.flush()
        return BulkReindexRead(
            total_document_count=len(documents),
            indexed_document_count=indexed_count,
            failed_document_count=len(failures),
            failures=failures,
        )

    async def rechunk_and_reindex_all(self, session: AsyncSession, user_id: UUID) -> BulkReindexRead:
        if not settings.embedding_base_url or not settings.embedding_api_key:
            raise AppError("EMBEDDING_NOT_CONFIGURED", "请先配置 Embedding 服务", status_code=409)
        documents = await self.repo.list_documents(session, user_id)
        failures: list[dict[str, str]] = []
        indexed_count = 0
        for document in documents:
            try:
                await self.rechunk_and_reindex(session, user_id, document.id)
                indexed_count += 1
            except AppError as exc:
                document.index_status = "KEYWORD_READY"
                failures.append({"document_id": str(document.id), "name": document.name, "reason": exc.message})
        await session.flush()
        return BulkReindexRead(
            total_document_count=len(documents),
            indexed_document_count=indexed_count,
            failed_document_count=len(failures),
            failures=failures,
        )

    async def search_with_trace(
        self, session: AsyncSession, user_id: UUID, query: str, scope: list[str], top_k: int
    ) -> tuple[list[KnowledgeSearchResult], dict[str, str | bool], int]:
        started_at = time.perf_counter()
        candidates = await self.repo.chunks_for_user(session, user_id, scope)
        fts_ranked = await self.repo.fts_chunks_for_user(session, user_id, scope, query)
        fallback_ranked = sorted(
            candidates,
            key=lambda pair: self._keyword_score(query, pair[0].content),
            reverse=True,
        )
        fallback_ranked = [
            pair for pair in fallback_ranked if self._keyword_score(query, pair[0].content) >= 0.12
        ][:20]
        # FTS is efficient for tokenized languages; character n-grams supplement
        # it for Chinese and phrasing/order changes. Keep FTS results first but
        # retain unique fallback candidates to avoid zero-recall surprises.
        keyword_ranked = fts_ranked + [
            pair for pair in fallback_ranked if pair[0].id not in {chunk.id for chunk, _ in fts_ranked}
        ]
        rankings = [[chunk.id for chunk, _ in keyword_ranked]]
        by_id = {chunk.id: (chunk, document) for chunk, document in candidates}
        lexical_scores = {chunk.id: self._keyword_score(query, chunk.content) for chunk, _ in candidates}
        semantic_scores: dict[UUID, float] = {}
        if settings.embedding_base_url and settings.embedding_api_key:
            vector = (await self.embeddings.embed([query]))[0]
            vector_ranked = await self.repo.vector_chunks_for_user(session, user_id, scope, vector)
            rankings.append([chunk.id for chunk, _, _ in vector_ranked])
            by_id.update({chunk.id: (chunk, document) for chunk, document, _ in vector_ranked})
            semantic_scores = {chunk.id: max(0.0, 1 - distance) for chunk, _, distance in vector_ranked}
        fused = rrf_fuse(rankings)
        max_lexical_score = max(lexical_scores.values(), default=1)
        max_fused_score = max(fused.values(), default=1)
        reranked = []
        for chunk_id, fused_score in fused.items():
            lexical_score = lexical_scores.get(chunk_id, 0) / max_lexical_score
            rrf_score = fused_score / max_fused_score
            if semantic_scores:
                score = 0.6 * semantic_scores.get(chunk_id, 0) + 0.25 * lexical_score + 0.15 * rrf_score
            else:
                score = 0.75 * lexical_score + 0.25 * rrf_score
            reranked.append((chunk_id, score))
        reranked.sort(key=lambda item: item[1], reverse=True)
        if reranked:
            # Avoid exposing weak tail candidates just because a small corpus has
            # fewer than the requested number of chunks.
            threshold = max(0.18, reranked[0][1] * 0.55)
            reranked = [item for item in reranked if item[1] >= threshold]
        selected: list[tuple[UUID, float]] = []
        per_document_count: dict[UUID, int] = defaultdict(int)
        for chunk_id, score in reranked:
            document_id = by_id[chunk_id][1].id
            if per_document_count[document_id] >= 2:
                continue
            selected.append((chunk_id, score))
            per_document_count[document_id] += 1
            if len(selected) >= top_k:
                break
        results = [
            KnowledgeSearchResult(
                document_id=by_id[chunk_id][1].id,
                chunk_id=chunk_id,
                content=by_id[chunk_id][0].content,
                source_type=by_id[chunk_id][1].source_type,
                source_name=by_id[chunk_id][1].name,
                score=score,
            )
            for chunk_id, score in selected
        ]
        retrieval_config: dict[str, str | bool] = {
            "lexical_retriever": "postgres_fts+char_ngram" if fts_ranked else "char_ngram_fallback",
            "vector_retriever": bool(settings.embedding_base_url and settings.embedding_api_key),
            "fusion": "rrf+hybrid_rerank",
        }
        latency_ms = round((time.perf_counter() - started_at) * 1000)
        return results, retrieval_config, latency_ms

    async def search(
        self, session: AsyncSession, user_id: UUID, query: str, scope: list[str], top_k: int
    ) -> list[KnowledgeSearchResult]:
        return (await self.search_with_trace(session, user_id, query, scope, top_k))[0]

    async def record_feedback(
        self,
        session: AsyncSession,
        user_id: UUID,
        search_event_id: UUID,
        chunk_id: UUID,
        relevance: str,
    ) -> None:
        event = await self.repo.get_search_event(session, user_id, search_event_id)
        if str(chunk_id) not in event.result_chunk_ids_json:
            raise AppError("VALIDATION_ERROR", "该引用不属于本次检索结果", status_code=422)
        await self.repo.upsert_feedback(
            session,
            KnowledgeSearchFeedback(
                user_id=user_id,
                search_event_id=event.id,
                chunk_id=chunk_id,
                relevance=relevance,
            ),
        )

    async def quality_overview(
        self, session: AsyncSession, user_id: UUID
    ) -> KnowledgeRetrievalQualityRead:
        events = await self.repo.list_search_events(session, user_id)
        feedback = await self.repo.list_feedback(session, user_id)
        search_count = len(events)
        feedback_count = len(feedback)
        return KnowledgeRetrievalQualityRead(
            search_count=search_count,
            zero_result_rate=(sum(event.result_count == 0 for event in events) / search_count) if search_count else 0,
            average_latency_ms=(sum(event.latency_ms for event in events) / search_count) if search_count else 0,
            feedback_count=feedback_count,
            feedback_coverage_rate=(feedback_count / sum(event.result_count for event in events)) if events else 0,
            helpful_rate=(sum(item.relevance == "helpful" for item in feedback) / feedback_count)
            if feedback_count
            else None,
        )

    @staticmethod
    def _keyword_score(query: str, content: str) -> float:
        query_tokens = _tokens(query)
        normalized = content.casefold()
        token_score = sum(normalized.count(token) for token in query_tokens)
        query_ngrams = _character_ngrams(query)
        content_ngrams = _character_ngrams(content)
        ngram_recall = len(query_ngrams & content_ngrams) / len(query_ngrams) if query_ngrams else 0
        return token_score + ngram_recall
